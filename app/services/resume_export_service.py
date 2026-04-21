from __future__ import annotations

import os
from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Any
from uuid import UUID

from sqlalchemy.orm import Session

from app.models.resume import Resume
from app.models.user import User
from app.services.resume_templates import build_resume_payload, get_template_spec, normalize_template_id, render_resume_html

BASE_DIR = Path(__file__).resolve().parents[1]


def _load_latest_resume(db: Session, user_id: UUID) -> Resume | None:
    return (
        db.query(Resume)
        .filter(Resume.user_id == user_id)
        .order_by(Resume.updated_at.desc())
        .first()
    )


def _resolve_resume_payload(db: Session, user_id: UUID, template_id: str | None = None) -> tuple[dict[str, Any], str]:
    user = db.query(User).filter(User.id == user_id).first()
    if not user:
        raise ValueError("User not found")

    resume = _load_latest_resume(db, user_id)
    if resume:
        resume_json = dict(resume.resume_json or {})
        selected_template = normalize_template_id(template_id or resume.selected_template)
        return resume_json, selected_template

    fallback_resume = {
        "personal": {
            "name": user.name or "CareerOS Candidate",
            "email": user.email or "",
            "summary": "",
            "links": [],
        },
        "experience": [],
        "projects": [],
        "education": [],
        "skills": [],
        "optional_sections": [],
    }
    return fallback_resume, normalize_template_id(template_id)


def _file_name_from_resume(resume_json: dict[str, Any], suffix: str) -> str:
    payload = build_resume_payload(resume_json)
    base = str(payload.get("name") or "resume").strip() or "resume"
    safe = "".join(char if char.isalnum() else "_" for char in base).strip("_") or "resume"
    return f"{safe}.{suffix}"


def export_resume_pdf(db: Session, user_id: UUID, template_name: str | None = None) -> tuple[str, str]:
    """Render selected template and export to PDF using WeasyPrint."""
    try:
        from weasyprint import HTML
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("WeasyPrint is not available. Install dependencies and system libraries.") from exc

    resume_json, selected_template = _resolve_resume_payload(db, user_id, template_name)
    html = render_resume_html(selected_template, resume_json)

    with NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        HTML(string=html).write_pdf(tmp.name)
        file_path = tmp.name

    download_name = _file_name_from_resume(resume_json, "pdf")
    return file_path, download_name


def export_resume_docx(db: Session, user_id: UUID, template_name: str | None = None) -> tuple[str, str]:
    """Export resume to DOCX using python-docx with template-aware section styling."""
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Pt
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("python-docx is not available. Install required dependencies.") from exc

    resume_json, selected_template = _resolve_resume_payload(db, user_id, template_name)
    template = get_template_spec(selected_template)
    payload = build_resume_payload(resume_json)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = template.font_family_body.split(",")[0].strip()
    normal_style.font.size = Pt(template.font_size_body)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER if template.header_style == "centered" else WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run(str(payload.get("name") or "CareerOS Candidate"))
    title_run.bold = True
    title_run.font.name = template.font_family_heading.split(",")[0].strip()
    title_run.font.size = Pt(template.font_size_name)

    if payload.get("title"):
        title_line = doc.add_paragraph()
        title_line.alignment = title.alignment
        title_line.add_run(str(payload["title"]))

    contact_parts = [value for value in [payload.get("email"), payload.get("phone"), payload.get("location")] if value]
    contact_parts.extend(payload.get("links", []))
    if contact_parts:
        contact = doc.add_paragraph()
        contact.alignment = title.alignment
        contact.add_run(" | ".join(str(item) for item in contact_parts if str(item).strip()))

    def add_paragraph_text(text: str) -> None:
        text = text.strip()
        if text:
            doc.add_paragraph(text)

    def add_section_heading(text: str) -> None:
        heading = doc.add_paragraph()
        run = heading.add_run(text)
        run.bold = True
        run.font.name = template.font_family_heading.split(",")[0].strip()
        run.font.size = Pt(template.font_size_heading)

    def add_entry(entry: dict[str, Any] | str, fields: tuple[str, ...]) -> None:
        if isinstance(entry, str):
            add_paragraph_text(entry)
            return

        heading_parts = [str(entry.get(field, "")).strip() for field in fields if str(entry.get(field, "")).strip()]
        if heading_parts:
            doc.add_paragraph(" · ".join(heading_parts))

        description = str(entry.get("description") or entry.get("summary") or "").strip()
        if description:
            add_paragraph_text(description)

        bullets = entry.get("bullets") if isinstance(entry.get("bullets"), list) else []
        for bullet in bullets:
            bullet_text = str(bullet).strip()
            if bullet_text:
                doc.add_paragraph(bullet_text, style="List Bullet")

    if payload.get("summary"):
        add_section_heading("Summary")
        add_paragraph_text(str(payload["summary"]))

    section_groups: list[tuple[str, list[Any], tuple[str, ...]]] = [
        ("Experience", payload.get("experience", []), ("title", "company", "duration")),
        ("Projects", payload.get("projects", []), ("title", "company", "duration")),
        ("Education", payload.get("education", []), ("institution", "degree", "year")),
    ]

    if template.layout != "two-column":
        section_groups.append(("Skills", payload.get("skills", []), ()))

    for section_title, items, fields in section_groups:
        if not items:
            continue
        add_section_heading(section_title)
        if section_title == "Skills":
            skills = [str(item).strip() for item in items if str(item).strip()]
            if template.skills_style == "tags":
                add_paragraph_text(", ".join(skills))
            elif template.skills_style == "grouped":
                add_paragraph_text(" · ".join(skills))
            else:
                for skill in skills:
                    add_paragraph_text(skill)
            continue

        for entry in items:
            add_entry(entry, fields)

    if template.layout == "two-column" and payload.get("skills"):
        add_section_heading("Skills")
        skills = [str(item).strip() for item in payload.get("skills", []) if str(item).strip()]
        if template.skills_style == "tags":
            add_paragraph_text(", ".join(skills))
        elif template.skills_style == "grouped":
            add_paragraph_text(" · ".join(skills))
        else:
            for skill in skills:
                add_paragraph_text(skill)

    for section_item in payload.get("optional_sections", []):
        title_text = str(section_item.get("title") or section_item.get("key") or "Additional Section").strip()
        items = section_item.get("items") if isinstance(section_item.get("items"), list) else []
        if not items:
            continue
        add_section_heading(title_text)
        for item in items:
            add_paragraph_text(str(item))

    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        doc.save(tmp.name)
        file_path = tmp.name

    download_name = _file_name_from_resume(resume_json, "docx")
    return file_path, download_name


def cleanup_export_file(path: str) -> None:
    """Delete temporary exported file after response has been sent."""
    try:
        if path and os.path.exists(path):
            os.remove(path)
    except OSError:
        pass
