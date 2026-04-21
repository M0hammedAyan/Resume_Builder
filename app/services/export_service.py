from __future__ import annotations

from pathlib import Path
from tempfile import NamedTemporaryFile
from typing import Any

from app.services.resume_templates import build_resume_payload, get_template_spec, normalize_template_id, render_resume_html


def _clean(value: Any) -> str:
    return str(value or "").strip()


def _section_items(resume_json: dict[str, Any], section: str) -> list[dict[str, Any] | str]:
    items = resume_json.get(section)
    if isinstance(items, list):
        return items
    return []


def _skills_text(resume_json: dict[str, Any]) -> str:
    skills = resume_json.get("skills")
    if isinstance(skills, list):
        values = [_clean(item) for item in skills if _clean(item)]
        return ", ".join(values)
    if isinstance(skills, str):
        return _clean(skills)
    return ""


def _file_name_from_resume(resume_json: dict[str, Any], suffix: str) -> str:
    payload = build_resume_payload(resume_json)
    base = _clean(payload.get("name")) or "resume"
    safe = "".join(char if char.isalnum() else "_" for char in base).strip("_") or "resume"
    return f"{safe}.{suffix}"


def generate_pdf(resume_json: dict[str, Any], template_id: str | None = None) -> tuple[str, str]:
    try:
        from weasyprint import HTML
    except Exception as exc:  # noqa: BLE001
        raise RuntimeError("WeasyPrint is required for PDF export") from exc

    selected_template = normalize_template_id(template_id or resume_json.get("selected_template") or None)
    html = render_resume_html(selected_template, resume_json)

    with NamedTemporaryFile(delete=False, suffix=".pdf") as tmp:
        path = Path(tmp.name)

    HTML(string=html).write_pdf(str(path))
    return str(path), _file_name_from_resume(resume_json, "pdf")


def generate_docx(resume_json: dict[str, Any], template_id: str | None = None) -> tuple[str, str]:
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Pt

    selected_template = normalize_template_id(template_id or resume_json.get("selected_template") or None)
    template = get_template_spec(selected_template)
    payload = build_resume_payload(resume_json)

    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    normal_style = doc.styles["Normal"]
    normal_style.font.name = template.font_family_body.split(",")[0].strip()
    normal_style.font.size = Pt(template.font_size_body)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER if template.header_style == "centered" else WD_ALIGN_PARAGRAPH.LEFT
    title_run = title.add_run(payload["name"])
    title_run.bold = True
    title_run.font.name = template.font_family_heading.split(",")[0].strip()
    title_run.font.size = Pt(template.font_size_name)

    if payload["title"]:
        title_line = doc.add_paragraph()
        title_line.alignment = title.alignment
        title_line.add_run(payload["title"])

    contact_parts = [value for value in [payload["email"], payload["phone"], payload["location"]] if value]
    contact_parts.extend(payload["links"])
    if contact_parts:
        contact = doc.add_paragraph()
        contact.alignment = title.alignment
        contact.add_run(" | ".join(contact_parts))

    def add_paragraph_text(text: str) -> None:
        if text.strip():
            doc.add_paragraph(text.strip())

    def add_section_heading(text: str) -> None:
        heading = doc.add_paragraph()
        heading_run = heading.add_run(text)
        heading_run.bold = True
        heading_run.font.name = template.font_family_heading.split(",")[0].strip()
        heading_run.font.size = Pt(template.font_size_heading)

    def add_entry(entry: dict[str, Any] | str, fields: tuple[str, ...]) -> None:
        if isinstance(entry, str):
            add_paragraph_text(entry)
            return

        heading_parts = [_clean(entry.get(field)) for field in fields if _clean(entry.get(field))]
        if heading_parts:
            doc.add_paragraph(" · ".join(heading_parts))

        description = _clean(entry.get("description") or entry.get("summary"))
        if description:
            add_paragraph_text(description)

        bullets = entry.get("bullets") if isinstance(entry.get("bullets"), list) else []
        for bullet in bullets:
            bullet_text = _clean(bullet)
            if bullet_text:
                doc.add_paragraph(bullet_text, style="List Bullet")

    if payload["summary"]:
        add_section_heading("Summary")
        add_paragraph_text(payload["summary"])

    if template.layout == "two-column":
        main_sections = [
            ("Experience", payload["experience"], ("title", "company", "duration")),
            ("Projects", payload["projects"], ("title", "company", "duration")),
        ]
        side_sections = [
            ("Education", payload["education"], ("institution", "degree", "year")),
            ("Skills", payload["skills"], ()),
        ]
    else:
        main_sections = [
            ("Experience", payload["experience"], ("title", "company", "duration")),
            ("Projects", payload["projects"], ("title", "company", "duration")),
            ("Education", payload["education"], ("institution", "degree", "year")),
            ("Skills", payload["skills"], ()),
        ]
        side_sections = []

    for title_text, items, fields in main_sections:
        if not items:
            continue
        add_section_heading(title_text)
        for entry in items:
            add_entry(entry, fields)

    for title_text, items, fields in side_sections:
        if not items:
            continue
        add_section_heading(title_text)
        if title_text == "Skills":
            if template.skills_style == "tags":
                add_paragraph_text(", ".join(items))
            elif template.skills_style == "grouped":
                add_paragraph_text(" · ".join(items))
            else:
                for skill in items:
                    add_paragraph_text(str(skill))
            continue
        for entry in items:
            add_entry(entry, fields)

    if payload["optional_sections"]:
        for section_item in payload["optional_sections"]:
            title_text = _clean(section_item.get("title")) or _clean(section_item.get("key")) or "Additional Section"
            items = section_item.get("items") if isinstance(section_item.get("items"), list) else []
            if not items:
                continue
            add_section_heading(title_text)
            for item in items:
                add_paragraph_text(_clean(item))

    with NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
        path = Path(tmp.name)
        doc.save(str(path))

    return str(path), _file_name_from_resume(resume_json, "docx")
