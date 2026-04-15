from __future__ import annotations

import io
import re
from dataclasses import dataclass
from pathlib import Path

from docx import Document
from pdfminer.high_level import extract_text as pdfminer_extract_text

try:
    import fitz
except Exception:  # noqa: BLE001
    fitz = None

try:
    import pytesseract
    from PIL import Image
except Exception:  # noqa: BLE001
    pytesseract = None
    Image = None


@dataclass
class ParsedDocument:
    text: str
    metadata: dict


class DocumentParser:
    """Parse resume files with OCR fallback when extraction is too sparse."""

    def __init__(self, min_text_length: int = 300) -> None:
        self.min_text_length = min_text_length

    def parse(self, file_content: bytes, filename: str) -> ParsedDocument:
        ext = Path(filename).suffix.lower()
        metadata = {
            "source_file": filename,
            "parser_used": None,
            "ocr_used": False,
            "pdf_has_tables": False,
            "pdf_has_images": False,
            "pdf_fonts": [],
        }

        if ext == ".pdf":
            text = self._parse_pdf(file_content, metadata)
        elif ext == ".docx":
            text = self._parse_docx(file_content)
            metadata["parser_used"] = "python-docx"
        elif ext == ".txt":
            text = file_content.decode("utf-8", errors="ignore")
            metadata["parser_used"] = "plain-text"
        else:
            raise ValueError(f"Unsupported file extension: {ext}")

        if self._needs_ocr(text) and ext == ".pdf":
            ocr_text = self._ocr_pdf(file_content)
            if len(ocr_text.strip()) > len(text.strip()):
                text = ocr_text
            metadata["ocr_used"] = bool(ocr_text.strip())

        return ParsedDocument(text=text.strip(), metadata=metadata)

    def _parse_pdf(self, file_content: bytes, metadata: dict) -> str:
        text = self._parse_pdf_pdfminer(file_content)
        metadata["parser_used"] = "pdfminer.six"
        self._inspect_pdf_layout(file_content, metadata)

        if self._needs_ocr(text):
            fallback_text = self._parse_pdf_pymupdf(file_content)
            if len(fallback_text.strip()) > len(text.strip()):
                text = fallback_text
                metadata["parser_used"] = "pymupdf"

        return text

    def _parse_pdf_pdfminer(self, file_content: bytes) -> str:
        with io.BytesIO(file_content) as buffer:
            return pdfminer_extract_text(buffer) or ""

    def _parse_pdf_pymupdf(self, file_content: bytes) -> str:
        if fitz is None:
            return ""
        text_parts: list[str] = []
        with fitz.open(stream=file_content, filetype="pdf") as doc:  # type: ignore[union-attr]
            for page in doc:
                text_parts.append(page.get_text("text"))
        return "\n".join(text_parts)

    def _parse_docx(self, file_content: bytes) -> str:
        document = Document(io.BytesIO(file_content))
        paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
        return "\n".join(paragraphs)

    def _ocr_pdf(self, file_content: bytes) -> str:
        if fitz is None or pytesseract is None or Image is None:
            return ""

        text_parts: list[str] = []
        with fitz.open(stream=file_content, filetype="pdf") as doc:  # type: ignore[union-attr]
            for page in doc:
                pix = page.get_pixmap(dpi=220)
                image = Image.open(io.BytesIO(pix.tobytes("png")))
                text_parts.append(pytesseract.image_to_string(image))
        return "\n".join(text_parts)

    def _inspect_pdf_layout(self, file_content: bytes, metadata: dict) -> None:
        if fitz is None:
            return

        fonts: set[str] = set()
        has_images = False
        has_tables = False

        with fitz.open(stream=file_content, filetype="pdf") as doc:  # type: ignore[union-attr]
            for page in doc:
                if page.get_images(full=True):
                    has_images = True

                # Basic table detection heuristic.
                if hasattr(page, "find_tables"):
                    try:
                        tables = page.find_tables()
                        if getattr(tables, "tables", []):
                            has_tables = True
                    except Exception:  # noqa: BLE001
                        pass

                for font_info in page.get_fonts(full=True):
                    if len(font_info) >= 4:
                        fonts.add(str(font_info[3]).lower())

        metadata["pdf_has_images"] = has_images
        metadata["pdf_has_tables"] = has_tables
        metadata["pdf_fonts"] = sorted(fonts)

    def _needs_ocr(self, text: str) -> bool:
        compact = re.sub(r"\s+", "", text or "")
        return len(compact) < self.min_text_length
